"""Extract plain text from PDF and Word (.docx) for resume / JD fields."""

from __future__ import annotations

import re
from io import BytesIO

from pypdf import PdfReader


def normalize_text(s: str) -> str:
    s = s.replace("\r\n", "\n").replace("\r", "\n")
    s = re.sub(r"[ \t\f\v]+", " ", s)
    s = re.sub(r"\n{3,}", "\n\n", s)
    return s.strip()


def extract_pdf_text(data: bytes) -> str:
    try:
        reader = PdfReader(BytesIO(data), strict=False)
    except Exception as e:
        raise ValueError(f"无法读取 PDF：{e}") from e
    if getattr(reader, "is_encrypted", False):
        try:
            unlocked = reader.decrypt("")
        except Exception as e:
            raise ValueError(f"PDF 解密失败：{e}") from e
        if unlocked == 0:
            raise ValueError("PDF 已加密且需要密码，请在本地解密后上传。")
    chunks: list[str] = []
    for page in reader.pages:
        try:
            t = page.extract_text() or ""
        except Exception as e:
            raise ValueError(f"读取 PDF 页面失败：{e}") from e
        if t.strip():
            chunks.append(t)
    return normalize_text("\n".join(chunks))


def extract_docx_text(data: bytes) -> str:
    try:
        from docx import Document
    except ImportError as e:
        raise ValueError("缺少 python-docx 依赖，请安装 requirements.txt。") from e
    try:
        doc = Document(BytesIO(data))
    except Exception as e:
        raise ValueError(f"无法解析 Word 文件（请确认是 .docx，不是旧版 .doc）：{e}") from e
    parts: list[str] = []
    for p in doc.paragraphs:
        t = (p.text or "").strip()
        if t:
            parts.append(t)
    for table in doc.tables:
        for row in table.rows:
            cells = [(c.text or "").strip() for c in row.cells]
            cells = [c for c in cells if c]
            if cells:
                parts.append(" | ".join(dict.fromkeys(cells)))
    return normalize_text("\n".join(parts))
